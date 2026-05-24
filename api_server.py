
import asyncio
import io
import json
import os
import sys
import tempfile
import threading
import uuid
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

import requests

PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT))

from fastapi import FastAPI, HTTPException, Query, UploadFile, File, BackgroundTasks, Depends
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from sse_starlette.sse import EventSourceResponse
from sqlalchemy.orm import Session

# Lealone风格核心组件
from agent.core.service_registry import service_registry
from agent.core.health_agent import health_agent
from agent.knowledge.kg_qa import kg_qa
from agent.router.router import health_router  # 智能路由器

# 服务实例
from agent.services.user_service import user_service
from agent.services.health_report_service import health_report_service
from agent.services.chat_history_service import chat_history_service
from agent.services.document_service import document_service
from rag.vector_store import VectorStoreService
from agent.tools.health_enums import (
    get_genders, get_activity_levels, get_health_goals,
    get_fitness_levels, get_exercise_types, get_diet_types
)
from agent.tools.health_tools import calculate_bmi_structured, calculate_daily_calorie_structured
from project.config_hander import chroma_conf
from project.file_hander import get_file_md5_hex
from project.path_tool import get_abs_path
from project.logger_handler import logger
from agent.database.db_config import init_db, check_db_connection, get_db
from agent.database.neo4j_config import neo4j_conn

_building_lock = threading.Lock()


@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期"""
    print("=" * 60)
    print("  健康智能助手 - Lealone风格")
    print("  访问 http://localhost:7958")
    print("  零代码对话式健康管理")
    print("=" * 60)

    # 初始化数据库
    init_db()
    print("[DB] 数据库初始化完成")

    # 验证 Neo4j 连接（非阻塞，失败仅警告）
    try:
        neo4j_conn.verify_connectivity()
        print("[Neo4j] 连接验证成功")
    except Exception as e:
        print(f"[Neo4j] 连接验证失败（图谱功能不可用）: {e}")

    # 注册服务处理器
    _register_service_handlers()

    yield

    # 清理
    neo4j_conn.close()
    print("健康智能助手已停止")


def _register_service_handlers():
    """注册服务处理器到服务注册中心"""
    # 对话服务处理器（仅 api_server 注册）
    service_registry.register_handler("chat_service", "get_history",
        lambda sid=None: chat_history_service.load_history(sid))
    service_registry.register_handler("chat_service", "save_history",
        lambda msgs, sid=None: chat_history_service.save_history(msgs, sid))
    service_registry.register_handler("chat_service", "clear_history",
        lambda sid=None: chat_history_service.clear_history(sid))


app = FastAPI(
    title="健康智能助手 API",
    version="2.0.0",
    description="Lealone风格零代码健康管理系统",
    lifespan=lifespan
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 静态文件
web_dir = PROJECT_ROOT / "web"
app.mount("/static", StaticFiles(directory=str(web_dir)), name="static")


# ========== Pydantic 模型 ==========

class UserCreate(BaseModel):
    name: str
    gender: str = "男"
    age: int = 25
    height: float = 170
    weight: float = 65
    activity_level: str = "轻度活动"
    health_goal: str = "保持健康"

class UserUpdate(BaseModel):
    name: Optional[str] = None
    gender: Optional[str] = None
    age: Optional[int] = None
    height: Optional[float] = None
    weight: Optional[float] = None
    activity_level: Optional[str] = None
    health_goal: Optional[str] = None

class HealthRecordAdd(BaseModel):
    date: str
    weight: Optional[float] = None
    blood_pressure_systolic: Optional[int] = None
    blood_pressure_diastolic: Optional[int] = None
    heart_rate: Optional[int] = None
    sleep_hours: Optional[float] = None
    sleep_quality: Optional[int] = None
    steps: Optional[int] = None
    calories_intake: Optional[float] = None

class ChatMessage(BaseModel):
    role: str
    content: str

class ServiceExecute(BaseModel):
    service: str
    method: str
    params: Dict = {}


# ========== Lealone风格服务架构API ==========

@app.get("/api/schema")
async def get_service_schema():
    """获取服务架构描述 - Lealone风格"""
    return health_agent.get_service_schema()


@app.post("/api/execute")
async def execute_service(data: ServiceExecute):

    try:
        result = health_agent.execute_service(data.service, data.method, data.params)
        return {"success": True, "result": result}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ========== 枚举 API ==========

@app.get("/api/v1/enums/all")
async def get_all_enums():
    return {
        "genders": get_genders(),
        "activity_levels": get_activity_levels(),
        "health_goals": get_health_goals(),
        "fitness_levels": get_fitness_levels(),
        "exercise_types": get_exercise_types(),
        "diet_types": get_diet_types(),
    }


# ========== 用户 API ==========

@app.get("/api/v1/users")
async def list_users():
    users = user_service.list_users()
    return {"users": users, "total": len(users)}

@app.get("/api/v1/users/{user_id}")
async def get_user(user_id: str):
    user = user_service.get_user(user_id)
    if user is None:
        raise HTTPException(status_code=404, detail=f"用户 {user_id} 不存在")
    return user

@app.post("/api/v1/users")
async def create_user(data: UserCreate):
    with _building_lock:
        result = user_service.create_user(**data.model_dump())
    return result

@app.put("/api/v1/users/{user_id}")
async def update_user(user_id: str, data: UserUpdate):
    with _building_lock:
        kwargs = {k: v for k, v in data.model_dump().items() if v is not None}
        if not kwargs:
            raise HTTPException(status_code=400, detail="没有提供要更新的字段")
        result = user_service.update_user(user_id, **kwargs)
    if result is None:
        raise HTTPException(status_code=404, detail=f"用户 {user_id} 不存在")
    return result

@app.delete("/api/v1/users/{user_id}")
async def delete_user(user_id: str):
    with _building_lock:
        success = user_service.delete_user(user_id)
    if not success:
        raise HTTPException(status_code=404, detail=f"用户 {user_id} 不存在")
    return {"message": f"用户 {user_id} 已删除"}

@app.post("/api/v1/users/{user_id}/health-record")
async def add_health_record(user_id: str, data: HealthRecordAdd):
    with _building_lock:
        success = user_service.add_health_record(user_id, **data.model_dump())
    if not success:
        raise HTTPException(status_code=404, detail=f"用户 {user_id} 不存在或数据添加失败")
    return {"message": "健康记录添加成功"}

@app.get("/api/v1/users/{user_id}/health-records")
async def get_health_records(user_id: str):
    user = user_service.get_user(user_id)
    if user is None:
        raise HTTPException(status_code=404, detail=f"用户 {user_id} 不存在")
    records = user.get("health_records", {})
    return {"user_id": user_id, "health_records": records, "total": len(records)}

@app.get("/api/v1/dashboard/stats")
async def get_dashboard_stats():
    """获取仪表盘统计数据"""
    # 用户统计
    users = user_service.list_users()
    user_count = len(users)

    # 健康记录统计
    health_record_count = 0
    for user_id, user in users.items():
        health_record_count += len(user.get("health_records", {}))

    # 对话统计
    chat_stats = chat_history_service.get_all_stats()

    # KG实体和关系统计
    kg_stats = health_kg.get_stats_detailed()

    # 知识库文档统计
    knowledge_stats = document_service.get_stats()

    return {
        "user_count": user_count,
        "health_record_count": health_record_count,
        "chat_session_count": chat_stats.get("session_count", 0),
        "chat_message_count": chat_stats.get("total_messages", 0),
        "kg_entity_count": kg_stats.get("total_entities", 0),
        "kg_relation_count": kg_stats.get("total_relations", 0),
        "knowledge_doc_count": knowledge_stats.get("total", 0),
        "knowledge_chunk_count": knowledge_stats.get("total_chunks", 0)
    }


@app.get("/api/v1/dashboard/user-demographics")
async def get_user_demographics():
    """获取用户群体统计数据"""
    users = user_service.list_users()

    # BMI分布统计
    bmi_categories = {"偏瘦": 0, "正常": 0, "超重": 0, "肥胖": 0}
    age_groups = {"18-25": 0, "26-35": 0, "36-45": 0, "46-55": 0, "56+": 0}
    gender_count = {"男": 0, "女": 0}
    activity_levels = {}

    for user_id, user in users.items():
        basic = user.get("basic_info", {})

        # BMI分类
        height = basic.get("height", 170)
        weight = basic.get("weight", 65)
        if height > 0 and weight > 0:
            bmi = weight / (height / 100) ** 2
            if bmi < 18.5:
                bmi_categories["偏瘦"] += 1
            elif bmi < 24:
                bmi_categories["正常"] += 1
            elif bmi < 28:
                bmi_categories["超重"] += 1
            else:
                bmi_categories["肥胖"] += 1

        # 年龄分组
        age = basic.get("age", 25)
        if age <= 25:
            age_groups["18-25"] += 1
        elif age <= 35:
            age_groups["26-35"] += 1
        elif age <= 45:
            age_groups["36-45"] += 1
        elif age <= 55:
            age_groups["46-55"] += 1
        else:
            age_groups["56+"] += 1

        # 性别统计
        gender = basic.get("gender", "男")
        gender_count[gender] = gender_count.get(gender, 0) + 1

        # 活动水平统计
        activity = basic.get("activity_level", "轻度活动")
        activity_levels[activity] = activity_levels.get(activity, 0) + 1

    return {
        "total_users": len(users),
        "bmi_distribution": bmi_categories,
        "age_distribution": age_groups,
        "gender_distribution": gender_count,
        "activity_distribution": activity_levels
    }


# ========== BMI / Calorie ==========

@app.get("/api/v1/health/bmi")
async def calculate_bmi_api(height: float, weight: float):
    if height <= 0 or weight <= 0:
        raise HTTPException(status_code=400, detail="身高和体重必须为正数")
    return calculate_bmi_structured(height, weight)

@app.get("/api/v1/health/daily-calorie")
async def calculate_daily_calorie_api(
    gender: str, age: int, height: float, weight: float,
    activity_level: str = "轻度活动"
):
    return calculate_daily_calorie_structured(gender, age, height, weight, activity_level)


# ========== 报告 API ==========

@app.get("/api/v1/reports")
async def list_reports():
    return {"reports": health_report_service.list_reports(), "total": len(health_report_service.list_reports())}

@app.get("/api/v1/reports/{report_id}")
async def get_report(report_id: str):
    report = health_report_service.get_report(report_id)
    if report is None:
        raise HTTPException(status_code=404, detail=f"报告 {report_id} 不存在")
    return report

@app.get("/api/v1/reports/search")
async def search_reports(q: str = Query(..., min_length=1)):
    return {"query": q, "results": health_report_service.search_reports(q)}


# ========== 聊天 API ==========

@app.get("/api/v1/chat/history")
async def load_chat_history(session_id: Optional[str] = None):
    messages = chat_history_service.load_history(session_id)
    info = chat_history_service.get_history_info(session_id)
    return {"messages": messages, "info": info}

@app.post("/api/v1/chat/history")
async def save_chat_history(data: List[ChatMessage], session_id: Optional[str] = None):
    messages = [{"role": m.role, "content": m.content} for m in data]
    success = chat_history_service.save_history(messages, session_id)
    return {"message": "聊天记录已保存" if success else "保存失败"}

@app.delete("/api/v1/chat/history")
async def clear_chat_history(session_id: Optional[str] = None):
    success = chat_history_service.clear_history(session_id)
    return {"message": "聊天记录已清空" if success else "清空失败"}

@app.get("/api/v1/chat/stream")
async def chat_stream(q: str = Query(..., min_length=1), session_id: Optional[str] = Query(None)):
    """SSE流式对话 - 智能路由分发，token级流式 + 中间步骤推送"""
    # 输入去噪
    from agent.preprocessing.input_denoiser import input_denoiser
    original_query = q
    denoised_query = input_denoiser.denoise(q) if input_denoiser else q.strip()

    # 从 ChatHistoryService 加载历史
    history = chat_history_service.load_recent_history(session_id)

    async def event_generator():
        try:
            async for event in health_router.route_stream_async(
                denoised_query, history=history, original_query=original_query
            ):
                yield event
        except Exception as e:
            yield {"event": "error", "data": str(e)}
            yield {"event": "done", "data": ""}

    return EventSourceResponse(event_generator())


# ========== 会话管理 API ==========

@app.get("/api/v1/chat/sessions")
async def list_sessions():
    """列出所有会话"""
    sessions = chat_history_service.list_sessions()
    return {"sessions": sessions, "total": len(sessions)}


@app.post("/api/v1/chat/sessions")
async def create_session():
    """新建会话，生成短 UUID 并预创建空历史文件"""
    short_id = uuid.uuid4().hex[:8]
    chat_history_service.save_history([], session_id=short_id)
    return {"session_id": short_id, "message": "会话创建成功"}


@app.delete("/api/v1/chat/sessions/{session_id}")
async def delete_session(session_id: str):
    """删除指定会话"""
    success = chat_history_service.delete_session(session_id)
    if not success:
        raise HTTPException(status_code=404, detail=f"会话 {session_id} 不存在或删除失败")
    return {"message": f"会话 {session_id} 已删除"}


@app.get("/api/v1/route/analyze")
async def analyze_route(q: str = Query(..., min_length=1)):
    """
    路由诊断API - 分析意图分类结果

    用于调试和查看路由决策过程
    """
    from agent.router.intent_classifier import intent_classifier
    result = intent_classifier.classify(q)
    return {
        "query": q,
        "intent": result.intent,
        "confidence": result.confidence,
        "is_complex": result.is_complex,
        "matched_pattern": result.matched_pattern
    }


# ========== 天气 API ==========

@app.get("/api/v1/weather")
async def get_weather_api(city: str = "北京"):
    try:
        url = f"https://wttr.in/{city}?format=j1&lang=zh"
        resp = await asyncio.to_thread(
            requests.get, url, timeout=10,
            headers={"User-Agent": "curl/7.68.0"},
        )
        data = resp.json()
        current = data.get("current_condition", [{}])[0]
        return {
            "city": city,
            "temp": current.get("temp_C", ""),
            "feels_like": current.get("FeelsLikeC", ""),
            "humidity": current.get("humidity", ""),
            "weather": current.get("lang_zh", [{}])[0].get("value", ""),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== 报告生成 ==========

@app.post("/api/v1/report/generate")
async def generate_health_report(user_ids: Optional[List[str]] = None):
    """生成健康分析Word报告"""
    if user_ids is None:
        user_ids = []

    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
        plt.rcParams['axes.unicode_minus'] = False
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="需要安装 python-docx 和 matplotlib")

    all_users = user_service.list_users()
    if user_ids:
        users_data = {uid: all_users[uid] for uid in user_ids if uid in all_users}
    else:
        users_data = all_users

    if not users_data:
        raise HTTPException(status_code=400, detail="没有可分析的用户数据")

    temp_dir = tempfile.mkdtemp()
    try:
        doc = Document()
        title = doc.add_heading("个人健康分析报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("注意：本报告仅供参考，不能替代专业医生的诊断和治疗建议。")

        # 用户概览
        doc.add_heading("一、用户概览", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = \
            "用户ID", "姓名", "性别/年龄", "身高/体重", "BMI"

        for uid, u in users_data.items():
            basic = u.get("basic_info", {})
            h = basic.get("height", 170)
            w = basic.get("weight", 65)
            bmi = w / (h / 100) ** 2
            row = table.add_row().cells
            row[0].text = uid
            row[1].text = basic.get("name", "未命名")
            row[2].text = f"{basic.get('gender', '')}/{basic.get('age', '')}岁"
            row[3].text = f"{h:.0f}cm/{w:.0f}kg"
            row[4].text = f"{bmi:.1f}"

        # AI建议
        doc.add_heading("二、健康改善建议", level=1)
        try:
            from model.factory import chat_model
            user_summary = []
            for uid, u in users_data.items():
                basic = u.get("basic_info", {})
                user_summary.append(f"- {basic.get('name', uid)}: BMI={basic.get('weight',65)/(basic.get('height',170)/100)**2:.1f}")

            prompt = f"请根据以下用户数据生成5条健康改善建议:\n{chr(10).join(user_summary)}"
            response = chat_model.invoke(prompt)
            ai_text = response.content if hasattr(response, 'content') else str(response)
            for line in ai_text.strip().split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
        except Exception:
            for fb in ["保持规律作息", "每周运动150分钟", "均衡饮食", "定期体检", "保持良好心态"]:
                doc.add_paragraph(f"- {fb}")

        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        from urllib.parse import quote
        filename = f"健康分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        encoded_filename = quote(filename)
        return StreamingResponse(
            doc_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"报告生成失败: {str(e)}")


# ========== 根路由 ==========

@app.get("/")
async def serve_index():
    return FileResponse(str(web_dir / "kg_index.html"))


# ========== 健康检查 ==========

@app.get("/health")
async def health_check():
    """数据库和 Neo4j 连接健康检查"""
    db_ok = check_db_connection()
    neo4j_ok = neo4j_conn.check_connection()
    return {
        "status": "ok" if db_ok and neo4j_ok else "degraded",
        "db": db_ok,
        "neo4j": neo4j_ok,
    }


# ========== 知识图谱 API ==========

from agent.knowledge.health_kg import health_kg
from agent.knowledge.kg_qa import kg_qa
from agent.knowledge.ner import health_ner
from agent.knowledge.ner import sync_ner_dictionary, update_entity_labels
from agent.knowledge.kg_extractor import kg_extractor

class KGQuery(BaseModel):
    query: str

class KGEntity(BaseModel):
    entity: str

class KGRelation(BaseModel):
    entity1: str
    entity2: str

@app.get("/api/v1/kg/schema")
async def get_kg_schema():
    """获取知识图谱架构"""
    return health_kg.get_schema()

@app.get("/api/v1/kg/stats")
async def get_kg_stats():
    """获取知识图谱统计信息"""
    stats = health_kg.get_stats_detailed()
    schema = health_kg.get_schema()
    return {
        "nodes": stats.get("total_entities", 0),
        "edges": stats.get("total_relations", 0),
        "entity_types": stats.get("entity_types", {}),
    }

@app.post("/api/v1/kg/query")
async def kg_query_api(data: KGQuery):
    """知识图谱问答"""
    result = kg_qa.answer(data.query)
    return result

@app.get("/api/v1/kg/entity/{entity_name}")
async def get_kg_entity(entity_name: str):
    """获取实体信息"""
    entity = health_kg.get_entity(entity_name)
    if not entity:
        raise HTTPException(status_code=404, detail=f"实体 '{entity_name}' 不存在")
    return {
        "entity": entity.to_dict(),
        "relations": health_kg.get_entity_relations(entity_name)
    }

@app.get("/api/v1/kg/entity/{entity_name}/relations")
async def get_entity_relations(entity_name: str):
    """获取实体关系"""
    relations = health_kg.get_entity_relations(entity_name)
    if not relations:
        raise HTTPException(status_code=404, detail=f"实体 '{entity_name}' 暂无关系数据")
    return {"entity": entity_name, "relations": relations}

@app.post("/api/v1/kg/relation")
async def query_entity_relation(data: KGRelation):
    """查询两个实体之间的关系"""
    paths = health_kg.find_path(data.entity1, data.entity2, max_depth=3)
    return {
        "entity1": data.entity1,
        "entity2": data.entity2,
        "paths": paths,
        "found": len(paths) > 0
    }

@app.get("/api/v1/kg/search")
async def search_kg_entities(q: str = Query(..., min_length=1)):
    """搜索实体"""
    results = health_kg.search_entities(q)
    return {"query": q, "results": results, "total": len(results)}

@app.post("/api/v1/kg/ner")
async def kg_ner_api(data: KGQuery):
    """命名实体识别"""
    results = health_ner.recognize(data.query)
    entities = [r for r in results if r["entity_type"] > 0]
    return {
        "text": data.query,
        "entities": entities,
        "annotated": health_ner.annotate_text(data.query)
    }

@app.get("/api/v1/kg/disease/{disease}/symptoms")
async def get_disease_symptoms(disease: str):
    """查询疾病症状"""
    symptoms = health_kg.find_entities_by_relation(disease, "具有症状", "out")
    return {"disease": disease, "symptoms": symptoms}

@app.get("/api/v1/kg/diseases")
async def get_all_diseases():
    """获取所有疾病列表"""
    diseases = health_kg.get_all_diseases()
    return {"diseases": diseases, "total": len(diseases)}

@app.get("/api/v1/kg/disease/{disease_name}/graph")
async def get_disease_graph(disease_name: str):
    """获取疾病的完整图谱数据"""
    graph_data = health_kg.get_disease_graph(disease_name)
    if "error" in graph_data:
        raise HTTPException(status_code=404, detail=f"疾病 '{disease_name}' 不存在")
    return graph_data

@app.get("/api/v1/kg/disease/{disease}/treatment")
async def get_disease_treatment(disease: str):
    """查询疾病治疗方法"""
    treatments = health_kg.find_entities_by_relation(disease, "治疗方式", "out")
    return {"disease": disease, "treatments": treatments}

@app.get("/api/v1/kg/food/{food}/nutrients")
async def get_food_nutrients(food: str):
    """查询食物营养成分"""
    nutrients_contains = health_kg.find_entities_by_relation(food, "含有", "out")
    nutrients_rich = health_kg.find_entities_by_relation(food, "富含", "out")
    attrs = health_kg.get_entity_attributes(food)
    return {
        "food": food,
        "nutrients_contains": nutrients_contains,
        "nutrients_rich": nutrients_rich,
        "attributes": attrs
    }


# ========== 知识库文档管理 API ==========

ALLOWED_FILE_TYPES = tuple(chroma_conf.get("allow_knowledge_file_type", ["txt", "pdf"]))
MAX_FILE_SIZE_MB = chroma_conf.get("max_file_size_mb", 10)
MAX_FILES_PER_UPLOAD = chroma_conf.get("max_files_per_upload", 5)

# 全局向量存储服务实例
_vector_store_service = None

def get_vector_store_service():
    """获取向量存储服务实例（延迟初始化）"""
    global _vector_store_service
    if _vector_store_service is None:
        _vector_store_service = VectorStoreService()
    return _vector_store_service


def process_document_background(doc_id: str, file_path: str):
    """后台处理文档的函数（包含知识图谱更新）"""
    try:
        # 更新状态为processing
        document_service.update_document_status(doc_id, "processing")

        # 1. 加载文档内容
        from langchain_community.document_loaders import PyPDFLoader, TextLoader
        documents = []
        if file_path.endswith(".txt"):
            documents = TextLoader(file_path, encoding='utf-8').load()
        elif file_path.endswith(".pdf"):
            documents = PyPDFLoader(file_path).load()

        if not documents:
            document_service.update_document_status(doc_id, "failed", error_message="文档内容为空")
            return

        # 合并文档内容
        full_text = "\n".join([doc.page_content for doc in documents])

        # 2. 实体关系抽取
        extraction_results = []
        try:
            # 分chunk抽取
            chunks = []
            for doc in documents:
                # 简单分片
                text = doc.page_content
                chunk_size = 500
                for i in range(0, len(text), chunk_size):
                    chunks.append(text[i:i+chunk_size])

            extraction_results = kg_extractor.extract_from_chunks(chunks, doc_id)
            logger.info(f"[文档处理] 抽取结果: {len(extraction_results)}个chunks有实体/关系")
        except Exception as e:
            logger.warning(f"[文档处理] 实体抽取失败: {e}")

        # 3. 合并抽取结果并更新知识图谱
        entity_count = 0
        relation_count = 0
        merged_result = None

        if extraction_results:
            merged_result = kg_extractor.merge_results(extraction_results)

            # 添加实体到图谱
            from agent.knowledge.entity_types import HealthEntity, EntityType
            for e in merged_result.entities:
                entity = HealthEntity(
                    name=e.get("name", ""),
                    entity_type=EntityType(e.get("entity_type", 16)),
                    description=e.get("description", ""),
                    source=e.get("source", ""),
                    attributes={"source_doc_id": doc_id}
                )
                health_kg.add_entity(entity)
                entity_count += 1

            # 添加关系到图谱
            from agent.knowledge.relation_types import HealthRelation
            for r in merged_result.relations:
                relation = HealthRelation(
                    entity1=r.get("entity1", ""),
                    entity1_type=health_kg.get_entity_type(r.get("entity1", "")) or 0,
                    relation=r.get("relation", ""),
                    entity2=r.get("entity2", ""),
                    entity2_type=health_kg.get_entity_type(r.get("entity2", "")) or 0,
                    source=doc_id
                )
                health_kg.add_relation(relation)
                relation_count += 1

            logger.info(f"[文档处理] 图谱更新: 实体{entity_count}个, 关系{relation_count}条")

        # 4. 同步NER词典
        try:
            if merged_result and merged_result.entities:
                new_entities = [(e.get("name", ""), e.get("entity_type", 0)) for e in merged_result.entities]
                update_entity_labels(new_entities)
        except Exception as e:
            logger.warning(f"[文档处理] NER词典同步失败: {e}")

        # 5. 向量存储（带图谱元数据）
        vs = get_vector_store_service()
        extraction_dict = merged_result.to_dict() if merged_result else None
        chunk_count = vs.load_single_document_with_kg(file_path, doc_id, extraction_dict)

        # 6. 更新状态为completed
        document_service.update_document_status(
            doc_id, "completed",
            chunk_count=chunk_count,
            entity_count=entity_count,
            relation_count=relation_count
        )

    except Exception as e:
        # 更新状态为failed
        document_service.update_document_status(doc_id, "failed", error_message=str(e))


@app.post("/api/v1/knowledge/upload")
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    """
    上传文档到知识库

    - 支持文件类型: txt, pdf
    - 最大文件大小: 10MB
    - 自动进行向量切分和存储
    - 异步后台处理
    """
    # 验证文件类型
    file_ext = file.filename.split(".")[-1].lower() if "." in file.filename else ""
    if file_ext not in ALLOWED_FILE_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型: {file_ext}. 支持的类型: {', '.join(ALLOWED_FILE_TYPES)}"
        )

    # 验证文件大小
    content = await file.read()
    file_size = len(content)
    max_size_bytes = MAX_FILE_SIZE_MB * 1024 * 1024

    if file_size > max_size_bytes:
        raise HTTPException(
            status_code=400,
            detail=f"文件大小超过限制: {file_size / (1024*1024):.2f}MB > {MAX_FILE_SIZE_MB}MB"
        )

    # 生成保存文件名（UUID前缀）
    doc_id = document_service.generate_doc_id()
    saved_filename = f"{doc_id}_{file.filename}"
    knowledge_dir = Path(get_abs_path("data/knowledge"))
    saved_path = knowledge_dir / saved_filename

    # 保存文件
    with open(saved_path, "wb") as f:
        f.write(content)

    # 计算MD5
    md5_hex = get_file_md5_hex(str(saved_path))

    # 检查重复
    duplicate_doc = document_service.check_file_duplicate(str(saved_path))
    if duplicate_doc:
        # 删除刚保存的文件
        os.remove(saved_path)
        raise HTTPException(
            status_code=400,
            detail=f"文件内容已存在: {duplicate_doc.get('original_filename', '未知文件')}"
        )

    # 创建文档记录
    record = document_service.create_document_record(
        original_filename=file.filename,
        saved_path=str(saved_path),
        md5_hex=md5_hex,
        file_size=file_size,
        doc_id=doc_id
    )

    # 添加后台处理任务
    background_tasks.add_task(process_document_background, doc_id, str(saved_path))

    return {
        "success": True,
        "doc_id": doc_id,
        "filename": file.filename,
        "file_size": file_size,
        "status": "processing",
        "message": "文件已上传，正在后台处理"
    }


@app.get("/api/v1/knowledge/documents")
async def list_knowledge_documents(
    status: Optional[str] = None,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100)
):
    """
    获取知识库文档列表

    - 支持状态过滤: pending, processing, completed, failed
    - 支持分页
    """
    result = document_service.list_documents(status=status, page=page, page_size=page_size)
    return result


@app.get("/api/v1/knowledge/documents/{doc_id}")
async def get_knowledge_document(doc_id: str):
    """获取单个文档详情"""
    doc = document_service.get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail=f"文档 {doc_id} 不存在")
    return doc


@app.delete("/api/v1/knowledge/documents/{doc_id}")
async def delete_knowledge_document(doc_id: str):
    """
    删除文档

    - 删除向量库中的chunks
    - 删除文件
    - 删除元数据记录
    """
    doc = document_service.get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail=f"文档 {doc_id} 不存在")

    # 从向量库删除
    try:
        vs = get_vector_store_service()
        vs.delete_by_doc_id(doc_id)
    except Exception as e:
        logger.error(f"删除向量库chunks失败: {e}")

    # 删除文档记录和文件
    success = document_service.delete_document(doc_id)

    if not success:
        raise HTTPException(status_code=500, detail="删除文档失败")

    return {"success": True, "message": f"文档 {doc_id} 已删除"}


@app.get("/api/v1/knowledge/status/{doc_id}")
async def get_document_status(doc_id: str):
    """获取文档处理状态"""
    doc = document_service.get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail=f"文档 {doc_id} 不存在")

    return {
        "doc_id": doc_id,
        "status": doc.get("status"),
        "chunk_count": doc.get("chunk_count", 0),
        "error_message": doc.get("error_message"),
        "processed_at": doc.get("processed_at")
    }


@app.get("/api/v1/knowledge/stats")
async def get_knowledge_stats():
    """获取知识库统计信息（含图谱实时数据）"""
    stats = document_service.get_stats()
    # 用知识图谱实时数据覆盖/补充实体和关系统计
    kg_stats = health_kg.get_stats_detailed()
    stats["total_entities"] = kg_stats["total_entities"]
    stats["total_relations"] = kg_stats["total_relations"]
    return stats


# ========== 知识图谱扩展 API ==========

@app.get("/api/v1/knowledge/documents/{doc_id}/kg-status")
async def get_document_kg_status(doc_id: str):
    """
    获取文档的图谱状态

    返回实体数量、关系数量等图谱相关信息
    """
    doc = document_service.get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail=f"文档 {doc_id} 不存在")

    return {
        "doc_id": doc_id,
        "status": doc.get("status"),
        "entity_count": doc.get("entity_count", 0),
        "relation_count": doc.get("relation_count", 0),
        "kg_processed": doc.get("entity_count", 0) > 0 or doc.get("relation_count", 0) > 0
    }


@app.get("/api/v1/kg/new-entities")
async def get_new_entities(limit: int = Query(100, ge=1, le=500)):
    """
    获取新增实体列表

    返回从文档中抽取的新增实体（非预定义实体）
    """
    entities = health_kg.get_new_entities(limit)
    return {
        "entities": entities,
        "total": len(entities),
        "limit": limit
    }


@app.post("/api/v1/kg/sync-ner")
async def sync_ner_api():
    """
    手动同步NER词典

    将知识图谱中的所有实体同步到NER词典
    """
    try:
        result = sync_ner_dictionary()
        return {
            "success": True,
            "synced_count": result.get("synced", 0),
            "total_count": result.get("total", 0),
            "message": f"NER词典同步成功，共{result.get('total', 0)}个实体"
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": "NER词典同步失败"
        }


@app.post("/api/v1/kg/save")
async def save_kg_api():
    """
    手动冷备份图谱到 JSON 文件

    将当前知识图谱导出到 JSON 文件（Neo4j 自动持久化，此为冷备份）
    """
    try:
        kg_data_path = get_abs_path("data/knowledge/kg_data.json")
        success = health_kg.export_to_json(kg_data_path)
        if success:
            stats = health_kg.get_stats_detailed()
            return {
                "success": True,
                "message": "知识图谱冷备份成功",
                "file_path": kg_data_path,
                "stats": stats
            }
        else:
            return {
                "success": False,
                "message": "知识图谱冷备份失败"
            }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": "知识图谱冷备份失败"
        }


@app.get("/api/v1/kg/extract-preview")
async def extract_preview_api(
    text: str = Query(..., min_length=10, max_length=2000)
):
    """
    预览抽取结果

    对输入文本进行实体关系抽取，返回预览结果（不实际添加到图谱）
    """
    try:
        result = kg_extractor.extract_from_text(text, "preview")
        return {
            "success": True,
            "text": text,
            "entities": result.entities,
            "relations": result.relations,
            "entity_count": len(result.entities),
            "relation_count": len(result.relations)
        }
    except Exception as e:
        return {
            "success": False,
            "text": text,
            "error": str(e),
            "entities": [],
            "relations": []
        }


@app.get("/api/v1/kg/stats-detailed")
async def get_kg_stats_detailed():
    """
    获取知识图谱详细统计信息

    包括预定义实体、新增实体、关系数量等详细统计
    """
    stats = health_kg.get_stats_detailed()
    return stats


@app.get("/api/v1/kg/overview-graph")
async def get_kg_overview_graph():
    """
    获取图谱概览可视化数据

    返回适合在概览页面展示的节点和边数据
    """
    try:
        # 获取所有疾病
        diseases = health_kg.get_all_diseases()

        nodes = []
        links = []
        processed_nodes = set()

        # 类别定义 - 与 health_kg.get_disease_graph 的 category_map 一致
        categories = [
            {"name": "疾病", "color": "#ef4444"},      # 红色  index=0
            {"name": "症状", "color": "#f97316"},      # 橙色  index=1
            {"name": "药物", "color": "#8b5cf6"},      # 紫色  index=2
            {"name": "食物", "color": "#22c55e"},      # 绿色  index=3
            {"name": "营养素", "color": "#0ea5e9"},    # 蓝色  index=4
            {"name": "运动", "color": "#14b8a6"},      # 青色  index=5
            {"name": "危险因素", "color": "#dc2626"},  # 深红色 index=6
            {"name": "饮食类型", "color": "#84cc16"},  # 淡绿色 index=7
            {"name": "健康目标", "color": "#f59e0b"},  # 黄/金色 index=8
            {"name": "人体部位", "color": "#eab308"},  # 黄色 index=9
            {"name": "生活习惯", "color": "#a855f7"},  # 紫色 index=10
            {"name": "治疗方法", "color": "#3b82f6"},  # 蓝色 index=11
            {"name": "医学检查", "color": "#6366f1"},  # 紫色 index=12
            {"name": "健康术语", "color": "#0284c7"},  # 蓝色 index=13
            {"name": "其他", "color": "#64748b"}       # 灰色 index=14
        ]

        # 选择主要疾病展示（最多5个）
        main_diseases = diseases[:5] if diseases else ["高血压", "糖尿病", "感冒"]

        for disease in main_diseases:
            graph_data = health_kg.get_disease_graph(disease)
            if "error" not in graph_data:
                for node in graph_data.get("nodes", []):
                    if node["name"] not in processed_nodes:
                        # 直接使用 get_disease_graph 返回的 category 值
                        nodes.append({
                            "name": node["name"],
                            "category": node.get("category", 14),
                            "symbolSize": 50 if node["name"] == disease else 30
                        })
                        processed_nodes.add(node["name"])

                for link in graph_data.get("links", []):
                    links.append({
                        "source": link["source"],
                        "target": link["target"],
                        "relation": link["relation"]
                    })

        return {
            "nodes": nodes,
            "links": links,
            "categories": categories,
            "total_nodes": len(nodes),
            "total_links": len(links)
        }

    except Exception as e:
        logger.error(f"获取图谱可视化数据失败: {e}")
        return {
            "nodes": [],
            "links": [],
            "categories": [],
            "error": str(e)
        }


@app.get("/api/v1/knowledge/search-with-kg")
async def search_with_kg_api(
    q: str = Query(..., min_length=1),
    k: int = Query(5, ge=1, le=20)
):
    """
    搜索向量库并返回图谱上下文

    结合向量检索和知识图谱信息
    """
    try:
        vs = get_vector_store_service()
        result = vs.search_with_kg_context(q, k)
        return {
            "success": True,
            "query": q,
            "chunks": result.get("chunks", []),
            "kg_entities": result.get("kg_entities", []),
            "kg_relations": result.get("kg_relations", []),
            "kg_context": result.get("kg_context", ""),
            "total_chunks": result.get("total_chunks", 0)
        }
    except Exception as e:
        return {
            "success": False,
            "query": q,
            "error": str(e),
            "chunks": [],
            "kg_entities": [],
            "kg_relations": [],
            "kg_context": ""
        }